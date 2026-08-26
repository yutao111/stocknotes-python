import struct
import zlib

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = "三日低吸_衰竭止跌与下影止跌说明.docx"
CHART = "三日低吸_K线结构图.png"


def write_png(path, width, height, pixels):
    rows = b"".join(b"\x00" + bytes(row) for row in pixels)
    def chunk(kind, data):
        return struct.pack(">I", len(data)) + kind + data + struct.pack(">I", zlib.crc32(kind + data) & 0xFFFFFFFF)
    with open(path, "wb") as file:
        file.write(b"\x89PNG\r\n\x1a\n")
        file.write(chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)))
        file.write(chunk(b"IDAT", zlib.compress(rows, 9)))
        file.write(chunk(b"IEND", b""))


def make_candle_chart():
    width, height = 1400, 720
    background = (255, 253, 248)
    pixels = [[*background] * width for _ in range(height)]

    def rect(x1, y1, x2, y2, color):
        for y in range(max(0, y1), min(height, y2 + 1)):
            for x in range(max(0, x1), min(width, x2 + 1)):
                pixels[y][x * 3:x * 3 + 3] = color

    def line(x1, y1, x2, y2, color, thickness=3):
        steps = max(abs(x2 - x1), abs(y2 - y1), 1)
        for step in range(steps + 1):
            x = round(x1 + (x2 - x1) * step / steps)
            y = round(y1 + (y2 - y1) * step / steps)
            rect(x - thickness // 2, y - thickness // 2, x + thickness // 2, y + thickness // 2, color)

    navy, green, red, gray, light, amber = (23, 59, 42), (49, 126, 91), (181, 82, 69), (100, 112, 105), (218, 229, 222), (201, 130, 62)
    rect(0, 0, width - 1, height - 1, background)
    rect(50, 45, 675, 665, (245, 249, 246))
    rect(725, 45, 1350, 665, (250, 247, 241))
    for left in (50, 725):
        rect(left, 45, left + 625, 48, navy)
        rect(left, 662, left + 625, 665, navy)
        rect(left, 45, left + 3, 665, navy)
        rect(left + 622, 45, left + 625, 665, navy)
        for y in (200, 360, 520):
            line(left + 35, y, left + 590, y, light, 2)

    # Left: exhaustion stop. Three prior candles and a very small signal-day body.
    left_x = [160, 290, 420, 550]
    prior = [(110, 220, 150, 200, green), (190, 350, 220, 300, red), (285, 455, 325, 405, red)]
    for x, (high, low, open_y, close_y, color) in zip(left_x[:3], prior):
        line(x, high, x, low, color, 5)
        rect(x - 28, min(open_y, close_y), x + 28, max(open_y, close_y), color)
    x = left_x[3]
    line(x, 300, x, 560, amber, 6)
    rect(x - 35, 408, x + 35, 420, amber)
    line(590, 414, 650, 414, navy, 2)
    line(590, 560, 650, 560, navy, 2)

    # Right: lower-shadow stop. Three prior candles and a clear lower shadow.
    right_x = [835, 965, 1095, 1225]
    for x, (high, low, open_y, close_y, color) in zip(right_x[:3], prior):
        line(x, high, x, low, color, 5)
        rect(x - 28, min(open_y, close_y), x + 28, max(open_y, close_y), color)
    x = right_x[3]
    line(x, 300, x, 600, amber, 6)
    rect(x - 35, 380, x + 35, 430, amber)
    line(1265, 405, 1330, 405, navy, 2)
    line(1265, 600, 1330, 600, navy, 2)
    write_png(CHART, width, height, pixels)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in kwargs[edge].items():
            element.set(qn("w:{}".format(key)), str(value))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for name, size, color in (("Title", 22, "173B2A"), ("Heading 1", 15, "173B2A"), ("Heading 2", 12, "315A78")):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text, bold=False, color=None, size=None, alignment=None):
    paragraph = doc.add_paragraph()
    if alignment:
        paragraph.alignment = alignment
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    return paragraph


def add_formula(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Consolas"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string("315A78")
    return paragraph


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, label in enumerate(headers):
        cell = header.cells[index]
        cell.text = label
        set_cell_shading(cell, "315A78")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row_index, row_data in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            cells[index].text = str(value)
            if row_index % 2 == 0:
                set_cell_shading(cells[index], "F2F7F4")
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    doc.add_paragraph()
    return table


def add_candle_chart(doc):
    doc.add_heading("K 线结构图", level=1)
    add_text(doc, "以下为示意图，比例用于理解形态结构，并非按真实行情比例绘制。绿色表示收红，深色表示收阴或小实体。", color="59645F")

    make_candle_chart()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(CHART, width=Cm(16.1))
    add_text(doc, "左图：信号日实体很小，重点是多空趋于平衡。右图：信号日留下长下影线，重点是盘中低点获得承接。", alignment=WD_ALIGN_PARAGRAPH.CENTER, color="59645F", size=9)
    doc.add_paragraph()

    add_text(doc, "更直观的价格案例", bold=True, color="173B2A")
    add_table(doc, ["项目", "衰竭止跌型示例", "下影止跌型示例"], [
        ["前一日收盘", "100.0", "100.0"],
        ["信号日开盘 / 最高 / 最低 / 收盘", "98.0 / 100.5 / 96.5 / 98.3", "99.0 / 100.0 / 95.0 / 98.5"],
        ["实体", "|98.3 - 98.0| = 0.3", "|98.5 - 99.0| = 0.5"],
        ["当日振幅", "100.5 - 96.5 = 4.0", "100.0 - 95.0 = 5.0"],
        ["实体占振幅", "0.3 / 4.0 = 7.5%，满足 <= 10%", "0.5 / 5.0 = 10%，满足 <= 20%"],
        ["下影线 / 实体", "不作为必需条件", "(98.5 - 95.0) / 0.5 = 7 倍，满足 >= 1 倍"],
    ], [3.4, 6.2, 6.2])


def build_document():
    doc = Document()
    style_document(doc)
    title = doc.add_heading("三日低吸：衰竭止跌型与下影止跌型", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(doc, "规则说明、K 线图示与计算示例", alignment=WD_ALIGN_PARAGRAPH.CENTER, color="59645F", size=11)
    add_text(doc, "适用范围：重点观察池的日线三日低吸提醒。该规则用于识别客观行情结构，不构成买卖建议或收益保证。", color="59645F")

    doc.add_heading("一、先满足共同前置条件", level=1)
    add_text(doc, "衰竭止跌型和下影止跌型都不是单独看一根 K 线，而是先检查连续 4 个交易日组成的窗口：")
    add_table(doc, ["交易日", "角色", "必须满足的条件"], [
        ["T-3", "基准日", "收红：收盘价 > 开盘价"],
        ["T-2", "前置下跌第 1 天", "收阴：收盘价 < 开盘价；收盘低于基准日"],
        ["T-1", "前置下跌第 2 天", "收阴：收盘价 < 开盘价；收盘继续低于 T-2"],
        ["T", "信号日", "在前低附近探低后，按不同止跌结构分类"],
    ], [2.1, 3.2, 10.5])
    add_formula(doc, "T-3 收盘 > T-2 收盘 > T-1 收盘")
    add_bullet(doc, "前置两日累计跌幅必须达到参数门槛，当前默认至少 8%。")
    add_bullet(doc, "信号日最低价可高于前置两日最低价，但当前默认最多高 4%。例如前低 100，允许的最高信号日低点为 104。")
    add_bullet(doc, "共同前置不通过时，不再进入“衰竭”或“下影”细分判断。")

    doc.add_page_break()
    add_candle_chart(doc)

    doc.add_heading("二、衰竭止跌型", level=1)
    add_text(doc, "核心判断：连续下跌后，信号日盘中仍会探低，但开盘与收盘贴近、实体很小，说明单日方向性减弱，卖压可能衰竭。它不是强势反转确认。")
    add_table(doc, ["细分条件", "当前默认值", "计算与含义"], [
        ["低点修复至少", "2.0%", "(信号日收盘 - 信号日最低) / 前一日收盘 >= 2.0%。衡量从盘中低点收回了多少昨日收盘价的百分点。"],
        ["实体占振幅最多", "10%", "|收盘 - 开盘| / (最高 - 最低) <= 10%。小实体表示开收盘接近。"],
        ["信号日最低涨跌", "-1.5%", "信号日收盘相对前一日收盘不能跌超 1.5%。"],
        ["信号日最高涨跌", "+3.0%", "信号日收盘相对前一日收盘不能涨超 3.0%，避免把强反转归为衰竭止跌。"],
        ["止跌型量比最多", "1.0", "收盘确认时：信号日成交量 / T-1 成交量 <= 1.0，即不放量。盘中候选暂不强制该项。"],
    ], [3.4, 2.4, 10.0])
    add_text(doc, "修复幅度为什么以“前一日收盘”作分母？", bold=True, color="173B2A")
    add_formula(doc, "修复幅度 = (信号日收盘 - 信号日最低) / 前一日收盘")
    add_text(doc, "此口径与日涨跌幅、前置累计跌幅共用昨日价格基准，便于不同价格水平的标的横向比较。它衡量的是实际收回的价格空间，而不是相对最低价被放大的百分比。")

    doc.add_heading("三、下影止跌型", level=1)
    add_text(doc, "核心判断：信号日盘中杀跌后出现明显承接，最低价被拉回，留下相对实体足够长的下影线。它代表低位承接出现，但单根下影线仍可能只是下跌过程中的短暂反抽。")
    add_table(doc, ["细分条件", "当前默认值", "计算与含义"], [
        ["低点修复至少", "1.5%", "(信号日收盘 - 信号日最低) / 前一日收盘 >= 1.5%。"],
        ["实体占振幅最多", "20%", "|收盘 - 开盘| / (最高 - 最低) <= 20%。允许的实体比衰竭止跌型更大。"],
        ["下影线 / 实体至少", "1.0 倍", "[min(开盘, 收盘) - 最低] / |收盘 - 开盘| >= 1。要求下影线至少和实体一样长。"],
        ["信号日涨跌范围", "-1.5% 至 +3.0%", "与衰竭止跌型相同，限制收盘既不能明显走弱，也不能反弹过强。"],
        ["通知阶段", "仅盘中候选", "默认不升级为收盘确认，需后续价格行为验证。"],
    ], [3.4, 2.4, 10.0])

    doc.add_heading("四、两种形态的关键区别", level=1)
    add_table(doc, ["维度", "衰竭止跌型", "下影止跌型"], [
        ["主要观察点", "小实体，单日方向性消失", "长下影，低位承接清晰"],
        ["实体占振幅上限", "10%，更严格", "20%，可略放宽"],
        ["下影线硬条件", "无", "下影线至少为实体 1 倍"],
        ["低点修复门槛", "2.0%", "1.5%"],
        ["提醒处理", "可盘中候选，也可收盘确认", "默认仅盘中候选"],
        ["风险理解", "可能只是下跌动能减弱", "可能只是盘中短暂承接"],
    ], [3.2, 6.3, 6.3])

    doc.add_heading("五、快速判断清单", level=1)
    for item in (
        "先看 4 日窗口：基准日收红，随后两日收阴且收盘逐日走低。",
        "再看位置：信号日低点是否仍在前低附近，并达到共同累计跌幅要求。",
        "若实体很小，优先按衰竭止跌型检查。",
        "若低点被明显拉回，且下影线至少与实体等长，按下影止跌型检查。",
        "形态命中只表示出现候选结构；仍应结合大盘、板块、量能、仓位与风险控制独立决策。",
    ):
        add_bullet(doc, item)

    doc.add_heading("附：当前默认参数速览", level=1)
    add_table(doc, ["参数", "衰竭止跌型", "下影止跌型"], [
        ["低点修复至少", "2.0%", "1.5%"],
        ["实体占振幅最多", "10%", "20%"],
        ["下影 / 实体至少", "不要求", "1.0 倍"],
        ["信号日涨跌范围", "-1.5% 至 +3.0%", "-1.5% 至 +3.0%"],
        ["收盘确认", "允许，且需量比 <= 1.0", "默认不允许"],
    ], [4.0, 5.6, 5.6])

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
